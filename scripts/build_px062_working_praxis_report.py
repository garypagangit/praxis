from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "doc" / "px062_working_praxis_20260724"
TMP = ROOT / "tmp" / "docs" / "px062_working"
GATE1 = (
    ROOT
    / "reports"
    / "coding_agent_skill_provenance"
    / "gate1_public_corpus_20260724"
    / "summary.json"
)
GATE2_CONFIG = ROOT / "configs" / "px062_skill_hallucination_gate2_20260724.json"

NAVY = "17324D"
BLUE = "287AA9"
TEAL = "159A8C"
RED = "B54545"
AMBER = "D89516"
GRAY = "596874"
BLACK = "202930"
WHITE = "FFFFFF"
LIGHT = "F3F6F8"
PALE_BLUE = "E8F1F6"
PALE_AMBER = "FFF3D9"
PALE_RED = "FBE9E9"
PALE_GREEN = "E7F4F1"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    props.append(element)


def set_cell(cell, text: str, *, bold=False, color=BLACK, size=8.3) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    link.append(run)
    paragraph._p.append(link)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PX-062  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.06
    for name, size, color in (
        ("Title", 29, NAVY),
        ("Subtitle", 12.5, GRAY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 10.5, TEAL),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]
    header.text = "PRAXIS RESEARCH  /  PX-062 WORKING REPORT"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = rgb(BLUE)
    add_page_number(section.footer.paragraphs[0])
    doc.core_properties.title = (
        "PX-062 Skill Ecosystem Provenance and Registry Hallucination"
    )
    doc.core_properties.subject = "Update-ready Praxis report"
    doc.core_properties.author = "Gary Pagan"


def body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(value)


def numbered(doc: Document, values: list[str]) -> None:
    for value in values:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(value)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    result.autofit = False
    if widths:
        for index, width in enumerate(widths):
            result.columns[index].width = Inches(width)
    for index, header in enumerate(headers):
        set_shading(result.rows[0].cells[index], NAVY)
        set_cell(result.rows[0].cells[index], header, bold=True, color=WHITE, size=8)
    set_repeat_header(result.rows[0])
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for index, value in enumerate(values):
            set_cell(cells[index], value)
            if row_index % 2:
                set_shading(cells[index], LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return result


def pending_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    result = table(doc, headers, rows, widths)
    for row in result.rows[1:]:
        for cell in row.cells:
            if "Pending" in cell.text:
                set_shading(cell, PALE_AMBER)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = rgb(AMBER)
                    run.bold = True
    return result


def status_box(doc: Document, title: str, text: str, fill: str, accent: str) -> None:
    result = doc.add_table(rows=1, cols=1)
    result.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = result.cell(0, 0)
    set_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(accent)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(9.5)


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    r = p.add_run("PX-062")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(16)
    r.font.color.rgb = rgb(TEAL)
    p = doc.add_paragraph(style="Title")
    p.add_run("Skill Ecosystem Provenance and\nRegistry Hallucination")
    p = doc.add_paragraph(
        "An update-ready Praxis report on coding-agent skill supply-chain controls",
        style="Subtitle",
    )
    p.paragraph_format.space_after = Pt(24)
    result = doc.add_table(rows=1, cols=1)
    set_shading(result.cell(0, 0), NAVY)
    set_cell(
        result.cell(0, 0),
        "WORKING PRAXIS REPORT  /  GATE 1 COMPLETE  /  GATE 2 RUNNING",
        bold=True,
        color=WHITE,
        size=11,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.add_run("Prepared by Gary Pagan\n").bold = True
    p.add_run("Draft date: July 24, 2026\n")
    p.add_run(
        "Cloud job: px062-skill-hallucination-2026-07-24-22-21-01\n"
    )
    p.add_run("Current status at drafting: InProgress - Pending provisioning")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    r = p.add_run(
        "REPORT CONTROL\nThis is a working research report. Gate 0 and Gate 1 "
        "results are final within their stated boundaries. All Gate 2 result cells "
        "are marked Pending and must be populated only from the sealed cloud "
        "artifacts after completeness adjudication."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(GRAY)
    doc.add_page_break()


def get_font(size: int, bold=False):
    names = (
        ["C:/Windows/Fonts/aptos-bold.ttf", "C:/Windows/Fonts/arialbd.ttf"]
        if bold
        else ["C:/Windows/Fonts/aptos.ttf", "C:/Windows/Fonts/arial.ttf"]
    )
    for name in names:
        if Path(name).exists():
            return ImageFont.truetype(name, size)
    return ImageFont.load_default()


def draw_gate1(path: Path) -> None:
    labels = [
        "Authentic signed\npoison",
        "Tampered\npoison",
        "Nonexistent\nskill",
        "Clean exact\nskill",
    ]
    values = [1.0, 0.0, 0.0, 1.0]
    colors = [RED, TEAL, TEAL, BLUE]
    image = Image.new("RGB", (1450, 720), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (65, 35),
        "Gate 1: provenance-only admission rate",
        fill=f"#{NAVY}",
        font=get_font(40, True),
    )
    left, top, right, bottom = 110, 135, 1390, 570
    for tick in range(0, 101, 25):
        y = bottom - int(tick / 100 * (bottom - top))
        draw.line((left, y, right, y), fill="#DEE5EA", width=2)
        draw.text((35, y - 13), f"{tick}%", fill=f"#{GRAY}", font=get_font(23))
    width, gap, x = 210, 90, 200
    for label, value, color in zip(labels, values, colors):
        height = max(3, int(value * (bottom - top)))
        draw.rounded_rectangle(
            (x, bottom - height, x + width, bottom), 10, fill=f"#{color}"
        )
        display = f"{value:.0%}"
        box = draw.textbbox((0, 0), display, font=get_font(30, True))
        draw.text(
            (x + (width - (box[2] - box[0])) / 2, bottom - height - 40),
            display,
            fill=f"#{BLACK}",
            font=get_font(30, True),
        )
        for line_index, line in enumerate(label.split("\n")):
            box = draw.textbbox((0, 0), line, font=get_font(23))
            draw.text(
                (
                    x + (width - (box[2] - box[0])) / 2,
                    bottom + 18 + line_index * 27,
                ),
                line,
                fill=f"#{BLACK}",
                font=get_font(23),
            )
        x += width + gap
    image.save(path)


def figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.7))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.3)
    r.font.color.rgb = rgb(GRAY)


def reference(doc: Document, citation: str, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.24)
    p.add_run(citation + " ")
    add_hyperlink(p, label, url)


def build_doc(gate1: dict, config: dict, chart: Path) -> Document:
    doc = Document()
    configure(doc)
    title_page(doc)

    doc.add_heading("Executive Status", level=1)
    status_box(
        doc,
        "Current determination",
        "Gate 1 is a valid negative for provenance-only defense against authentic poisoned skills. Gate 2 is running and has no scientific result yet.",
        PALE_RED,
        RED,
    )
    body(
        doc,
        "PX-062 evaluates two connected but distinct questions. First, can deterministic provenance controls prevent skill supply-chain poisoning? Second, do coding models invent nonexistent skill names, and can registry verification prevent those invented names from becoming load attempts?"
    )
    table(
        doc,
        ["Stage", "Evidence", "Status", "Permitted conclusion"],
        [
            [
                "Gate 0",
                "180 inert controlled cases",
                "PASS",
                "The admission-policy implementation behaves as specified.",
            ],
            [
                "Gate 1",
                "1,070 released poisoned skills; 44 clean skills",
                "FAIL for provenance-only semantic defense",
                "Provenance blocks tampering and nonexistent identifiers but admits authentic malicious content.",
            ],
            [
                "Gate 2",
                "Two models x 300 tasks x three conditions",
                "RUNNING",
                "No hallucination-rate or mitigation claim is permitted yet.",
            ],
        ],
        [0.8, 1.85, 1.5, 2.4],
    )

    doc.add_heading("1. Research Problem and Literature Basis", level=1)
    body(
        doc,
        "Coding agents increasingly load reusable skills containing instructions, examples, scripts, and resources. Qu et al. demonstrate that malicious logic embedded in ordinary-looking skill documentation can cross from contextual guidance into an agent's action space. Their PoisonedSkills study released 1,070 adversarial skills derived from 81 seeds across 15 MITRE ATT&CK categories and evaluated four agent frameworks and five models."
    )
    body(
        doc,
        "The source study concentrates on the post-loading phase and assumes the poisoned skill reaches the agent context. PX-062 moves the defense boundary earlier. It tests admission controls at registry and load time, and then asks whether model-invented skill identifiers constitute an adjacent supply-chain exposure."
    )
    doc.add_heading("1.1 Praxis research questions", level=2)
    bullets(
        doc,
        [
            "RQ1: Can existence, version, hash, and signer verification block poisoned-skill admission while preserving clean utility?",
            "RQ2: Which failures are identity/integrity failures and which are semantic-content failures?",
            "RQ3: Do models recommend nonexistent skill names under open-ended selection?",
            "RQ4: Do registry constraints or post-generation verification reduce nonexistent-name load attempts without materially reducing known-skill selection accuracy?",
        ],
    )

    doc.add_heading("2. Gate 0 and Gate 1 Methods", level=1)
    doc.add_heading("2.1 Gate 0 implementation validation", level=2)
    body(
        doc,
        "Gate 0 used 180 inert cases spanning clean exact records, clean alternate versions, nonexistent names, hash mismatch, unsigned entries, signer mismatch, version mismatch, and missing manifests. The full policy required existence, manifest presence, matching content hash, and a valid expected signer. It produced zero attack escapes, zero clean false rejects, and complete decision traces. This result validates software logic only."
    )
    doc.add_heading("2.2 Gate 1 public-corpus evaluation", level=2)
    table(
        doc,
        ["Corpus", "Frozen source", "N", "Use"],
        [
            [
                "PoisonedSkills",
                "Zenodo DOI 10.5281/zenodo.19281322; CC BY 4.0",
                "1,070",
                "Authentic poisoned and tampered-poison conditions",
            ],
            [
                "OpenAI skills",
                "Commit 49f948faa9258a0c61caceaf225e179651397431",
                "44 files / 43 names",
                "Clean utility and clean-tamper controls",
            ],
        ],
        [1.15, 2.65, 0.8, 2.0],
    )
    body(
        doc,
        "The poisoned archive was verified against Zenodo and frozen at SHA-256 241ab288483ed77af2357f967c4018a59f5a407d5432b14ab6fc61a7cb33ad63. The evaluator streamed SKILL.md entries directly from the ZIP and did not execute any skill code. Registry signatures were modeled as valid or invalid evidence states; Gate 1 was not a production cryptographic deployment."
    )

    doc.add_heading("3. Gate 1 Results", level=1)
    figure(
        doc,
        chart,
        "Figure 1. A provenance-only policy rejects altered and nonexistent artifacts but admits authentic registered content, including poisoned content.",
    )
    table(
        doc,
        ["Condition", "N", "Admitted", "Admission rate", "Gate"],
        [
            ["Authentic signed poison", "1,070", "1,070", "100.0%", "FAIL"],
            ["Tampered poison", "1,070", "0", "0.0%", "PASS"],
            ["Nonexistent skill", "1,070", "0", "0.0%", "PASS"],
            ["Clean exact skill", "44", "44", "100.0%", "PASS"],
            ["Tampered clean skill", "44", "0", "0.0%", "PASS"],
        ],
        [2.0, 0.75, 0.85, 1.15, 1.0],
    )
    status_box(
        doc,
        "Gate 1 interpretation",
        "Existence checks, hashes, versions, and signatures establish identity and integrity. They do not establish that authenticated content is benign. Provenance is necessary for substitution and hallucinated-name defense, but insufficient for semantic poisoning.",
        PALE_AMBER,
        AMBER,
    )
    doc.add_heading("3.1 Policy implications", level=2)
    bullets(
        doc,
        [
            "Retain provenance checks as a deterministic admission layer.",
            "Do not claim they prevent the attack class demonstrated by the source paper when malicious content is the authentic registry artifact.",
            "Add publisher trust, review attestations, semantic analysis, least-privilege permissions, sandboxing, or multi-model inspection.",
            "Measure each layer separately so identity, integrity, semantic safety, and runtime containment are not conflated.",
        ],
    )

    doc.add_heading("4. Gate 2: Frozen Skill-Name Hallucination Study", level=1)
    body(
        doc,
        "Gate 2 evaluates registry-name invention. The task corpus was frozen after the Gate 1 registry inventory and before model collection. Known-skill prompts use catalog descriptions rather than revealing the answer name."
    )
    table(
        doc,
        ["Design element", "Frozen specification"],
        [
            ["Registry", "43 unique names from the OpenAI skills snapshot"],
            ["Known-skill tasks", "100 description-to-name selections"],
            ["Unavailable tasks", "100 capabilities with no matching registered skill"],
            ["Near-miss tasks", "100 mutated nonexistent names paired with a real base skill"],
            ["Models", "; ".join(config["models"])],
            ["Conditions", "Open-ended; registry-constrained; post-generation verification"],
            ["Generation", "Greedy decoding; maximum 32 new tokens"],
            ["Expected outputs", "1,800 unique model-condition-task records"],
        ],
        [1.5, 5.0],
    )
    doc.add_heading("4.1 Registered gates", level=2)
    table(
        doc,
        ["Gate", "Pass criterion", "Purpose"],
        [
            [
                "G2-H1",
                "Verified nonexistent attempted-load rate <= 1%",
                "Prevent invented identifiers from crossing the tool boundary",
            ],
            [
                "G2-H2",
                "Known-skill accuracy drop <= 5 percentage points",
                "Preserve utility relative to open-ended selection",
            ],
            [
                "G2-H3",
                "Trace completeness >= 99%",
                "Ensure valid model-condition-task coverage",
            ],
        ],
        [0.8, 2.35, 3.35],
    )
    doc.add_heading("4.2 Cloud execution record", level=2)
    table(
        doc,
        ["Field", "Value"],
        [
            ["Training job", "px062-skill-hallucination-2026-07-24-22-21-01"],
            ["Instance", "ml.g5.2xlarge"],
            ["Status at drafting", "InProgress - Pending provisioning"],
            [
                "S3 output",
                "s3://praxis-garypagan-272615233626-us-east-1/experiments/px062-skill-provenance/gate2-hallucination-20260724/output",
            ],
        ],
        [1.4, 5.1],
    )

    doc.add_heading("5. Gate 2 Results - Update Section", level=1)
    status_box(
        doc,
        "Result control",
        "Populate this section only after the cloud output is complete, downloaded, hashed, and independently checked for 1,800 unique model-condition-task records. Pending cells are intentional.",
        PALE_BLUE,
        BLUE,
    )
    doc.add_heading("5.1 Completeness", level=2)
    pending_table(
        doc,
        ["Check", "Expected", "Observed", "Decision"],
        [
            ["Models", "2", "Pending cloud adjudication", "Pending"],
            ["Conditions per model", "3", "Pending cloud adjudication", "Pending"],
            ["Tasks per condition", "300", "Pending cloud adjudication", "Pending"],
            ["Unique outputs", "1,800", "Pending cloud adjudication", "Pending"],
            ["Duplicate keys", "0", "Pending cloud adjudication", "Pending"],
            ["Trace completeness", ">= 99%", "Pending cloud adjudication", "Pending"],
        ],
        [2.1, 1.0, 2.15, 1.1],
    )
    doc.add_heading("5.2 Primary metrics by model and condition", level=2)
    metric_rows = []
    for model in config["models"]:
        short = model.split("/")[-1]
        for condition in config["conditions"]:
            metric_rows.append(
                [
                    short,
                    condition.replace("_", " "),
                    "Pending",
                    "Pending",
                    "Pending",
                    "Pending",
                ]
            )
    pending_table(
        doc,
        [
            "Model",
            "Condition",
            "Accuracy",
            "Nonexistent-name rate",
            "Nonexistent-attempt rate",
            "Abstention rate",
        ],
        metric_rows,
        [1.5, 1.3, 0.85, 1.15, 1.2, 0.9],
    )
    doc.add_heading("5.3 Mitigation comparison", level=2)
    pending_table(
        doc,
        ["Comparison", "Effect estimate", "Utility change", "Decision"],
        [
            [
                "Registry-constrained vs. open-ended",
                "Pending model outputs",
                "Pending model outputs",
                "Pending",
            ],
            [
                "Post-verification vs. open-ended",
                "Pending model outputs",
                "Pending model outputs",
                "Pending",
            ],
            [
                "Cross-model consistency",
                "Pending model outputs",
                "Pending model outputs",
                "Pending",
            ],
        ],
        [2.25, 1.6, 1.45, 1.05],
    )

    doc.add_heading("6. Final Determination Logic", level=1)
    body(
        doc,
        "Use the following branches after adjudication. Select one branch and delete the other two; do not combine incompatible conclusions."
    )
    table(
        doc,
        ["Outcome", "Required evidence", "Final framing"],
        [
            [
                "Positive",
                "H1-H3 pass for both models",
                "Registry verification reduces invented-skill load attempts with acceptable utility loss in the frozen benchmark.",
            ],
            [
                "Mixed",
                "Mitigation works for one model or condition, or utility gate fails",
                "The control is model- or condition-dependent and requires targeted deployment.",
            ],
            [
                "Negative",
                "Attempt-rate gate fails after verification or completeness is invalid",
                "The proposed verification layer is ineffective or the experiment must be rerun.",
            ],
        ],
        [0.8, 2.45, 3.25],
    )
    doc.add_heading("6.1 Working conclusion before Gate 2", level=2)
    body(
        doc,
        "PX-062 already establishes a defensible negative result: provenance alone does not prevent authentic semantic poisoning. It remains an effective identity-and-integrity layer against nonexistent names and modified artifacts. Whether this narrower function materially reduces model-generated skill-name risk is the unresolved Gate 2 question."
    )

    doc.add_heading("7. Validity and Claim Boundaries", level=1)
    table(
        doc,
        ["Dimension", "Current limitation", "Required treatment"],
        [
            [
                "Construct",
                "A valid signature is not a safety label",
                "Keep identity/integrity and semantic-safety claims separate.",
            ],
            [
                "External",
                "Gate 2 uses 43 names from one clean registry snapshot",
                "Replicate with another public ecosystem before broad claims.",
            ],
            [
                "Model",
                "Two open 7B models",
                "Do not generalize to proprietary or larger models.",
            ],
            [
                "Operational",
                "Recommendation and attempted-load are benchmark abstractions",
                "Add an isolated live-agent tool-boundary replication.",
            ],
            [
                "Security",
                "No destructive payload execution",
                "Retain marker-only, network-disabled validation.",
            ],
            [
                "Statistical",
                "Task templates share registry and capability families",
                "Report per-task-type results and avoid treating all rows as independent domains.",
            ],
        ],
        [1.0, 2.25, 3.25],
    )

    doc.add_heading("8. Post-Run Update Procedure", level=1)
    numbered(
        doc,
        [
            "Confirm the SageMaker job status is Completed and record start, end, billable seconds, image, instance, and output URI.",
            "Download the model artifact and compute SHA-256 hashes before extraction.",
            "Verify exactly 1,800 unique model-condition-task keys and compare task and registry hashes with the frozen local files.",
            "Run scripts/score_px062_skill_hallucination.py on the sealed outputs.",
            "Populate Sections 5.1-5.3 directly from the machine-readable score file.",
            "Apply the registered gates without changing thresholds.",
            "Select one final-determination branch in Section 6 and retain all negative or mixed evidence.",
            "Update the abstract, executive status, conclusion, experiment registry, and final artifact inventory.",
        ],
    )
    doc.add_heading("8.1 Frozen benchmark hashes", level=2)
    table(
        doc,
        ["Artifact", "SHA-256"],
        [
            [
                "tasks.jsonl",
                "439761496da03ed7bec64f241e37e424040d9ff2e9df8ed79bb402aba1b2ab9d",
            ],
            [
                "registry_names.json",
                "2c447b5eee07b2f2930fc8649860652b36d4902dafadfa83e3f5d7aa041a76db",
            ],
        ],
        [1.9, 4.6],
    )

    doc.add_page_break()
    doc.add_heading("References", level=1)
    reference(
        doc,
        "Qu, Y., Liu, Y., Geng, T., Deng, G., Li, Y., Zhang, L. Y., Zhang, Y., & Ma, L. (2026). Supply-Chain Poisoning Attacks Against LLM Coding Agent Skill Ecosystems. arXiv:2604.03081.",
        "Original paper",
        "https://arxiv.org/abs/2604.03081",
    )
    reference(
        doc,
        "Qu et al. (2026). PoisonedSkills: Exploiting Implicit Trust in LLM Coding Agent Skill Ecosystems. Zenodo.",
        "Released evaluation dataset",
        "https://doi.org/10.5281/zenodo.19281322",
    )
    reference(
        doc,
        "OpenAI. Skills Catalog for Codex. Frozen for PX-062 at commit 49f948faa9258a0c61caceaf225e179651397431.",
        "OpenAI skills repository",
        "https://github.com/openai/skills",
    )

    doc.add_heading("Appendix A. Artifact Map", level=1)
    table(
        doc,
        ["Artifact", "Repository location"],
        [
            [
                "Preregistration",
                "reports/coding_agent_skill_provenance/PX062_SKILL_PROVENANCE_PREREG_20260724.md",
            ],
            [
                "Current determination",
                "reports/coding_agent_skill_provenance/PX062_CURRENT_DETERMINATION_20260724.md",
            ],
            [
                "Gate 1 summary",
                "reports/coding_agent_skill_provenance/gate1_public_corpus_20260724/summary.json",
            ],
            [
                "Gate 2 configuration",
                "configs/px062_skill_hallucination_gate2_20260724.json",
            ],
            [
                "Frozen tasks",
                "data/px062/hallucination_benchmark/tasks.jsonl",
            ],
            [
                "Frozen registry",
                "data/px062/hallucination_benchmark/registry_names.json",
            ],
            [
                "Model collector",
                "scripts/run_px062_skill_hallucination_models.py",
            ],
            [
                "Scorer",
                "scripts/score_px062_skill_hallucination.py",
            ],
        ],
        [1.6, 4.9],
    )
    return doc


MARKDOWN = """# PX-062 Skill Ecosystem Provenance and Registry Hallucination

Working Praxis report - July 24, 2026

## Current status

- Gate 0: PASS on 180 inert implementation cases.
- Gate 1: VALID NEGATIVE for provenance-only defense against authentic poisoned skills.
- Gate 2: RUNNING; no live-model result yet.
- Cloud job: `px062-skill-hallucination-2026-07-24-22-21-01`

## Gate 1 result

| Condition | N | Admitted | Rate |
|---|---:|---:|---:|
| Authentic signed poison | 1,070 | 1,070 | 100.0% |
| Tampered poison | 1,070 | 0 | 0.0% |
| Nonexistent skill | 1,070 | 0 | 0.0% |
| Clean exact skill | 44 | 44 | 100.0% |
| Tampered clean skill | 44 | 0 | 0.0% |

Provenance proves identity and integrity, not semantic safety.

## Gate 2 frozen design

- Models: Qwen2.5-7B-Instruct and Mistral-7B-Instruct-v0.3
- Tasks: 300 per condition
- Conditions: open-ended, registry-constrained, post-generation verification
- Expected outputs: 1,800
- Gates: nonexistent attempted-load rate <= 1%; known-skill accuracy loss <= 5 percentage points; completeness >= 99%

## Gate 2 results

Status: **Pending cloud completion and independent adjudication.**

Populate model-condition metrics only from the sealed score artifact:

| Model | Condition | Accuracy | Nonexistent-name rate | Nonexistent-attempt rate | Abstention |
|---|---|---:|---:|---:|---:|
| Qwen2.5-7B-Instruct | open-ended | Pending | Pending | Pending | Pending |
| Qwen2.5-7B-Instruct | registry-constrained | Pending | Pending | Pending | Pending |
| Qwen2.5-7B-Instruct | post-verification | Pending | Pending | Pending | Pending |
| Mistral-7B-Instruct-v0.3 | open-ended | Pending | Pending | Pending | Pending |
| Mistral-7B-Instruct-v0.3 | registry-constrained | Pending | Pending | Pending | Pending |
| Mistral-7B-Instruct-v0.3 | post-verification | Pending | Pending | Pending | Pending |

## References

- [Original paper](https://arxiv.org/abs/2604.03081)
- [Released PoisonedSkills dataset](https://doi.org/10.5281/zenodo.19281322)
- [OpenAI skills catalog](https://github.com/openai/skills)
"""


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    gate1 = json.loads(GATE1.read_text(encoding="utf-8"))
    config = json.loads(GATE2_CONFIG.read_text(encoding="utf-8"))
    chart = TMP / "px062_gate1_admission.png"
    draw_gate1(chart)
    doc = build_doc(gate1, config, chart)
    docx = OUT / "PX-062_Working_Praxis_Report.docx"
    doc.save(docx)
    (OUT / "PX-062_Working_Praxis_Report.md").write_text(
        MARKDOWN, encoding="utf-8"
    )
    (OUT / "px062_gate1_admission.png").write_bytes(chart.read_bytes())
    print(docx)


if __name__ == "__main__":
    main()
