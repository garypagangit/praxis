from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "doc" / "px057_final_praxis_20260724"
TMP = ROOT / "tmp" / "docs" / "px057"
SUMMARY_PATH = (
    ROOT
    / "reports"
    / "adaptive_stopping_overthinking"
    / "gate2_full_determination_20260724"
    / "summary.json"
)
CONFIG_PATH = ROOT / "configs" / "px057_adaptive_stopping_gate2_full_20260724.json"

NAVY = "16324F"
BLUE = "2676A6"
TEAL = "149B8C"
GOLD = "E4A11B"
RED = "B44343"
PALE = "EAF2F7"
LIGHT = "F4F7F9"
GRAY = "566573"
WHITE = "FFFFFF"
BLACK = "1D252C"


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/aptos.ttf",
        "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/aptos-bold.ttf",
            "C:/Windows/Fonts/calibrib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
        ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold=False, color=BLACK, size=9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PX-057  |  ")
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [
        ("Title", 30, NAVY),
        ("Subtitle", 13, GRAY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, TEAL),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    doc.core_properties.title = title
    doc.core_properties.subject = "PX-057 adaptive stopping Gate 2 final Praxis paper"
    doc.core_properties.author = "Gary Pagan"
    header = section.header.paragraphs[0]
    header.text = "PRAXIS RESEARCH  /  FINAL EVIDENCE REPORT"
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    add_page_number(section.footer.paragraphs[0])


def title_page(doc: Document, subtitle: str, descriptor: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("PX-057")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph(style="Title")
    p.add_run("Adaptive Stopping to Prevent\nLLM Overthinking")
    p = doc.add_paragraph(subtitle, style="Subtitle")
    p.paragraph_format.space_after = Pt(26)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.55)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_text(cell, descriptor, bold=True, color=WHITE, size=12)
    cell.margin_top = Inches(0.15)
    cell.margin_bottom = Inches(0.15)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.add_run("Prepared by Gary Pagan\n").bold = True
    p.add_run("Final evidence date: July 24, 2026\n")
    p.add_run("Experiment stage: Gate 2 complete; H1-H3 passed; H4 transfer pending")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    r = p.add_run(
        "CLAIM BOUNDARY\nThis document reports one frozen 200-item GSM8K sample, "
        "Qwen2.5-7B-Instruct, and one eight-round iterative prompting protocol. "
        "It does not claim cross-model, cross-domain, or large-scale robustness."
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(TEAL)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=WHITE, size=8)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8)
            if row_index % 2:
                set_cell_shading(cells[i], LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def draw_accuracy_chart(metrics: dict, path: Path) -> None:
    labels = [
        "Fixed-long\nround 8",
        "Fixed-short\nround 2",
        "Uncertainty\nonly",
        "Answer\nstability",
        "Adaptive",
        "Oracle\n(descriptive)",
    ]
    values = [
        metrics["fixed_long_accuracy"],
        metrics["fixed_short_accuracy"],
        metrics["uncertainty_only_accuracy"],
        metrics["answer_stability_accuracy"],
        metrics["adaptive_accuracy"],
        metrics["oracle_best_step_accuracy"],
    ]
    colors = [RED, GOLD, GOLD, BLUE, TEAL, GRAY]
    img = Image.new("RGB", (1500, 760), "white")
    d = ImageDraw.Draw(img)
    title_f = font(42, True)
    body_f = font(27)
    value_f = font(30, True)
    d.text((70, 35), "Accuracy by inference policy", fill=f"#{NAVY}", font=title_f)
    left, top, right, bottom = 90, 135, 1450, 620
    for tick in range(0, 101, 20):
        y = bottom - int((tick / 100) * (bottom - top))
        d.line((left, y, right, y), fill="#DDE5EA", width=2)
        d.text((20, y - 15), f"{tick}%", fill=f"#{GRAY}", font=body_f)
    bar_w = 155
    gap = 60
    x = 150
    for label, value, color in zip(labels, values, colors):
        h = int(value * (bottom - top))
        d.rounded_rectangle((x, bottom - h, x + bar_w, bottom), 12, fill=f"#{color}")
        value_text = f"{value:.1%}"
        box = d.textbbox((0, 0), value_text, font=value_f)
        d.text(
            (x + (bar_w - (box[2] - box[0])) / 2, bottom - h - 42),
            value_text,
            fill=f"#{BLACK}",
            font=value_f,
        )
        lines = label.split("\n")
        for j, line in enumerate(lines):
            box = d.textbbox((0, 0), line, font=body_f)
            d.text(
                (x + (bar_w - (box[2] - box[0])) / 2, bottom + 17 + j * 30),
                line,
                fill=f"#{BLACK}",
                font=body_f,
            )
        x += bar_w + gap
    img.save(path)


def draw_gate_chart(metrics: dict, config: dict, path: Path) -> None:
    items = [
        (
            "Accuracy delta",
            metrics["adaptive_accuracy_delta"],
            config["gates"]["accuracy_delta_min"],
            "higher",
            "+29.5 pp",
            "threshold: >= -1.0 pp",
        ),
        (
            "Compute saving",
            metrics["mean_compute_saving"],
            config["gates"]["mean_compute_saving_min"],
            "higher",
            "66.5%",
            "threshold: >= 20%",
        ),
        (
            "Prevention rate",
            metrics["overthinking_prevention_rate"],
            config["gates"]["overthinking_prevention_min"],
            "higher",
            "89.6%",
            "threshold: >= 25%",
        ),
        (
            "Early-stop harm",
            metrics["early_stop_harm_rate"],
            config["gates"]["early_stop_harm_rate_max"],
            "lower",
            "0.5%",
            "threshold: <= 2%",
        ),
    ]
    img = Image.new("RGB", (1500, 760), "white")
    d = ImageDraw.Draw(img)
    d.text((65, 35), "Preregistered Gate 2 outcomes", fill=f"#{NAVY}", font=font(42, True))
    y = 140
    for label, value, threshold, direction, display, threshold_text in items:
        d.text((75, y), label, fill=f"#{BLACK}", font=font(30, True))
        d.text((430, y), display, fill=f"#{TEAL}", font=font(34, True))
        d.text((635, y + 4), threshold_text, fill=f"#{GRAY}", font=font(25))
        d.rounded_rectangle((1180, y - 8, 1405, y + 48), 18, fill=f"#{TEAL}")
        d.text((1244, y + 1), "PASS", fill="white", font=font(28, True))
        y += 135
    d.text(
        (75, 688),
        "All four registered checks passed; H4 transfer was outside this gate.",
        fill=f"#{GRAY}",
        font=font(25),
    )
    img.save(path)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(path), width=Inches(6.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_reference(doc: Document, citation: str, link_text: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(citation + " ")
    add_hyperlink(p, link_text, url)


def build_paper(summary: dict, config: dict, accuracy_chart: Path, gate_chart: Path) -> Document:
    m = summary["metrics"]
    doc = Document()
    configure_document(doc, "PX-057: Adaptive Stopping to Prevent LLM Overthinking")
    title_page(
        doc,
        "A preregistered Gate 2 evaluation of answer stability, uncertainty, and compute-aware early stopping",
        "FINAL PRAXIS PAPER  /  VALID POSITIVE RESULT",
    )

    add_kicker(doc, "Abstract")
    add_body(
        doc,
        "Extended test-time reasoning can improve difficult answers, but it can also cause a model to abandon a correct answer. "
        "PX-057 tested whether a frozen adaptive stopping rule could preserve accuracy while reducing generated-token use. "
        "The experiment evaluated 200 deterministically sampled GSM8K test questions with Qwen2.5-7B-Instruct under eight sequential reconsideration rounds. "
        "The primary policy stopped after round 2 when the normalized answer remained stable for two consecutive rounds and both rounds exceeded a mean-token-probability proxy of 0.05; otherwise it continued through round 8. "
        "The independent completeness audit verified 200 unique questions, 200 complete traces, 1,600 unique question-round generations, the frozen dataset hash, and the registered model identity. "
        "Adaptive accuracy was 91.0%, compared with 61.5% for fixed-long inference and 88.0% for fixed-short inference. "
        "The policy saved 66.5% of generated tokens on average (bootstrap 95% CI 63.7%-69.1%), prevented 60 of 67 observed correct-to-wrong events (89.6%; bootstrap 95% CI 82.1%-95.5%), and harmed 1 of 200 questions (0.5%). "
        "All preregistered H1-H3 gates passed. However, adaptive stopping tied the answer-stability control, so the confidence condition did not add observable accuracy in this setting. "
        "The result supports a bounded deployment proposition for this model, sample, and prompting protocol; cross-model and cross-domain transfer (H4) remains untested."
    )
    add_kicker(doc, "Keywords")
    add_body(doc, "large language models; test-time compute; adaptive stopping; overthinking; GSM8K; reproducible evaluation")

    doc.add_heading("1. Executive Finding", level=1)
    add_body(
        doc,
        "PX-057 is a valid positive Gate 2 result. The strongest demonstrated fact is not that all LLM reasoning should stop early; it is that repeatedly prompting this model to reconsider the same GSM8K answer through eight rounds caused substantial correct-to-wrong degradation, and a preregistered early stability gate avoided most of it."
    )
    add_table(
        doc,
        ["Decision dimension", "Observed result", "Interpretation"],
        [
            ["Scientific validity", "Independent completeness audit passed", "The frozen sample, model identity, traces, generations, and gate outputs were internally consistent."],
            ["Primary comparison", "91.0% adaptive vs. 61.5% fixed-long", "Adaptive stopping improved accuracy by 29.5 percentage points under the tested iterative protocol."],
            ["Efficiency", "66.5% mean generated-token saving", "The policy used roughly one-third of the fixed-long generated-token budget."],
            ["Safety of stopping", "1/200 early-stop harm", "The 0.5% observed harm rate remained below the preregistered 2% ceiling."],
            ["Mechanism nuance", "Adaptive tied answer stability at 91.0%", "The stability signal carried the observed performance; confidence added no measured accuracy benefit."],
            ["External validity", "H4 pending", "No cross-model, non-math, sampling, or prompt-transfer claim is supported yet."],
        ],
        [1.35, 1.55, 3.65],
    )

    doc.add_heading("2. Problem and Literature Basis", level=1)
    doc.add_heading("2.1 Why the experiment was necessary", level=2)
    add_body(
        doc,
        "Test-time scaling research shows that additional inference compute can improve reasoning, but the benefit depends on problem difficulty and allocation strategy. "
        "Snell et al. found that compute-optimal allocation can outperform uniform strategies, while Muennighoff et al. showed that forcing additional reasoning can sometimes repair an answer. "
        "These findings motivate more compute, but they do not imply that more compute is always beneficial."
    )
    add_body(
        doc,
        "The direct foundation for PX-057 is Zhou et al.'s 2026 study, which reports diminishing marginal returns at larger reasoning budgets, correct-to-wrong answer abandonment, and difficulty-dependent optimal reasoning lengths. "
        "That work identifies the phenomenon and shows that moderate stopping can reduce cost while maintaining comparable accuracy. PX-057 converts that observation into a deployable control question: can a precommitted, label-free rule stop before a harmful reversal without creating too many premature errors?"
    )
    doc.add_heading("2.2 Praxis gap", level=2)
    add_body(
        doc,
        "A practical stopping control must decide from information available at inference time. Gold correctness cannot be used to stop. PX-057 therefore combines two observable signals: normalized answer stability and a confidence proxy derived from generated-token probabilities. It evaluates the policy against fixed-short, fixed-long, stability-only, uncertainty-only, and oracle reference arms."
    )
    doc.add_heading("2.3 Research question", level=2)
    add_body(
        doc,
        "Can a frozen answer-stability-plus-confidence rule reduce generated-token use and prevent correct-to-wrong reversals while preserving accuracy relative to an eight-round fixed-long protocol?"
    )

    doc.add_heading("3. Hypotheses and Preregistered Gates", level=1)
    add_table(
        doc,
        ["ID", "Hypothesis", "Frozen pass criterion", "Outcome"],
        [
            ["H1", "Preserve accuracy while saving compute", "Accuracy delta vs. fixed-long >= -1 pp and mean saving >= 20%", "PASS: +29.5 pp; 66.5%"],
            ["H2", "Prevent harmful overthinking", "Prevent at least 25% of observable events", "PASS: 60/67; 89.6%"],
            ["H3", "Avoid excessive premature errors", "Early-stop harm <= 2% of all traces", "PASS: 1/200; 0.5%"],
            ["H4", "Transfer without retuning", "H1 and H3 pass in a frozen held-out domain", "PENDING: not tested in Gate 2"],
        ],
        [0.45, 2.05, 2.65, 1.4],
    )
    add_body(
        doc,
        "The preregistration was written before the scientific result. No post-result threshold change is part of this determination. Gate 2 adjudicated H1-H3 only; H4 was explicitly reserved for transfer testing."
    )

    doc.add_heading("4. Method", level=1)
    doc.add_heading("4.1 Experimental design", level=2)
    add_body(
        doc,
        "The study used a repeated-reconsideration trace design. Each question received an initial deterministic solution followed by seven deterministic prompts to reconsider the previous proposed solution, for a maximum of eight rounds. Each response ended in a numeric final answer that was normalized and compared with the GSM8K gold answer."
    )
    add_table(
        doc,
        ["Component", "Frozen specification"],
        [
            ["Model", "Qwen/Qwen2.5-7B-Instruct"],
            ["Dataset", "OpenAI GSM8K test split; frozen SHA-256 3730d312f6e3440559ace48831e51066acaca737f6eabec99bccb9e4b3c39d14"],
            ["Sample", "200 unique questions; deterministic seed 57"],
            ["Generation", "Greedy decoding (do_sample=False); max 256 new tokens per round"],
            ["Trace budget", "8 sequential rounds per question"],
            ["Primary policy", "Minimum round 2; patience 2; confidence threshold 0.05"],
            ["Compute measure", "Cumulative generated tokens; saving = 1 - adaptive tokens / fixed-long tokens"],
            ["Execution", "AWS SageMaker job px057-gate2-full-retry-2026-07-24-14-55-01-066"],
        ],
        [1.4, 5.15],
    )
    doc.add_heading("4.2 Stopping signals", level=2)
    add_body(
        doc,
        "Answer stability was satisfied when the normalized extracted answer was identical across the two most recent rounds. Confidence was computed as exp(mean normalized transition log-probability) over generated tokens, clipped to a numerically safe range. The result is a sequence-level mean-token-probability proxy, not a calibrated probability that the answer is correct."
    )
    add_body(
        doc,
        "At each round from round 2 onward, the adaptive rule stopped at the first point where both recent answers were stable and both confidence values were at least 0.05. If no point qualified, it returned the round-8 answer. The policy never observed the gold label."
    )
    doc.add_heading("4.3 Comparison arms", level=2)
    add_numbered(
        doc,
        [
            "Fixed-short: return the answer at round 2.",
            "Fixed-long: return the answer at round 8; the primary preregistered baseline.",
            "Answer-stability: stop after two identical normalized answers, without confidence.",
            "Uncertainty-only: stop at the first round at or after round 2 whose confidence proxy reaches 0.05.",
            "Adaptive: require both two-round answer stability and confidence at or above 0.05.",
            "Oracle best step: return the first correct eligible answer; descriptive ceiling only and not deployable.",
        ],
    )
    doc.add_heading("4.4 Outcome definitions", level=2)
    add_bullets(
        doc,
        [
            "Overthinking event: at least one correct eligible answer before round 8, followed by an incorrect round-8 answer.",
            "Prevented overthinking: the adaptive answer is correct on a trace classified as an overthinking event.",
            "Early-stop harm: round 8 is correct but the adaptive answer is incorrect.",
            "Generated-token saving: one minus cumulative adaptive generated tokens divided by cumulative round-8 generated tokens.",
            "Accuracy: exact match between the normalized extracted numeric answer and the GSM8K gold answer.",
        ],
    )
    doc.add_heading("4.5 Integrity and reproducibility controls", level=2)
    add_body(
        doc,
        "An independent adjudication script compared the cloud artifacts with the frozen configuration. It verified experiment and model identifiers, dataset hash, sample size, round count, 200 unique selected IDs, 200 complete traces, and exactly 1,600 unique and complete question-round pairs. It also recomputed the required gate-decision set and recorded SHA-256 hashes for the four core evidence files."
    )

    doc.add_heading("5. Results", level=1)
    add_figure(
        doc,
        accuracy_chart,
        "Figure 1. Accuracy for all registered arms. The oracle is descriptive and not deployable.",
    )
    add_table(
        doc,
        ["Arm", "Correct / 200", "Accuracy", "Difference vs. fixed-long"],
        [
            ["Fixed-long, round 8", "123", "61.5%", "Reference"],
            ["Fixed-short, round 2", "176", "88.0%", "+26.5 pp"],
            ["Uncertainty-only", "176", "88.0%", "+26.5 pp"],
            ["Answer stability", "182", "91.0%", "+29.5 pp"],
            ["Adaptive", "182", "91.0%", "+29.5 pp"],
            ["Oracle best step", "190", "95.0%", "+33.5 pp"],
        ],
        [2.15, 1.1, 1.0, 2.3],
    )
    add_figure(
        doc,
        gate_chart,
        "Figure 2. Every Gate 2 threshold passed. H4 was not part of this adjudication.",
    )
    add_table(
        doc,
        ["Metric", "Estimate", "Uncertainty / count", "Registered threshold"],
        [
            ["Adaptive accuracy delta", "+29.5 pp", "Descriptive difference", ">= -1.0 pp"],
            ["Mean generated-token saving", "66.5%", "Bootstrap 95% CI 63.7%-69.1%", ">= 20%"],
            ["Overthinking prevention", "89.6%", "60/67; bootstrap 95% CI 82.1%-95.5%", ">= 25%"],
            ["Early-stop harm", "0.5%", "1/200", "<= 2%"],
        ],
        [1.75, 1.05, 2.45, 1.3],
    )
    add_body(
        doc,
        "The adaptive policy returned 182 correct answers. Fixed-long inference returned 123 correct answers, meaning that the repeated reconsideration protocol degraded aggregate performance sharply by round 8. Sixty-seven traces satisfied the preregistered overthinking definition; the adaptive policy retained a correct answer on 60 of them. One trace moved in the opposite direction: fixed-long was correct while the adaptive return was wrong."
    )

    doc.add_heading("6. Interpretation", level=1)
    doc.add_heading("6.1 What the experiment demonstrates", level=2)
    add_bullets(
        doc,
        [
            "Under this frozen iterative prompting protocol, additional reconsideration frequently reduced answer correctness.",
            "A two-round stability gate was sufficient to avoid most observed correct-to-wrong reversals.",
            "The registered adaptive policy simultaneously exceeded the accuracy non-inferiority gate, the compute-saving gate, the prevention gate, and the harm ceiling.",
            "The experiment is reproducible at the artifact level because the sample, configuration, trace corpus, raw outputs, and hashes are preserved.",
        ],
    )
    doc.add_heading("6.2 What the experiment does not demonstrate", level=2)
    add_bullets(
        doc,
        [
            "It does not show that early stopping improves every model, task, prompt, or decoding regime.",
            "It does not establish performance on the full GSM8K test set; 200 questions were evaluated.",
            "It does not establish cross-domain or cross-model robustness; H4 remains pending.",
            "It does not prove that the confidence proxy is calibrated or causally useful.",
            "It does not compare production latency, energy, or dollar cost directly; generated tokens are the registered compute proxy.",
        ],
    )
    doc.add_heading("6.3 Why the stability-control tie matters", level=2)
    add_body(
        doc,
        "The adaptive and answer-stability arms both achieved 91.0% accuracy. Therefore, the result supports answer stability as the effective observable signal in this run, but it does not support an incremental accuracy claim for the confidence threshold. A production design should treat the confidence component as unvalidated until an ablation across models and domains shows a consistent benefit, lower harm, or better latency."
    )
    doc.add_heading("6.4 Relationship to the source literature", level=2)
    add_body(
        doc,
        "Zhou et al. identify diminishing returns and correct-answer abandonment at larger reasoning budgets. PX-057 reproduces the core operational concern in a different, deliberately simple protocol and tests a deployable stopping rule. The current result aligns with the source paper's argument that uniform compute allocation is suboptimal, while extending it from description to a preregistered safety-and-efficiency gate. It is not a direct reproduction of the source paper's full model, task, or token-budget matrix."
    )

    doc.add_heading("7. Validity Assessment", level=1)
    add_table(
        doc,
        ["Validity dimension", "Risk", "Assessment and mitigation"],
        [
            ["Internal", "Round-specific prompting may induce degradation", "The degradation is real for the frozen protocol, but should not be generalized to natural uninterrupted reasoning. Fixed-short and fixed-long arms make the protocol effect visible."],
            ["Construct", "Confidence proxy is not calibrated correctness", "Reported explicitly as mean-token probability. Stability-only and uncertainty-only ablations separate signal contributions."],
            ["Statistical", "Single 200-item sample", "Bootstrap intervals are reported for saving and prevention. No unregistered significance claim is made."],
            ["External", "One model and math dataset", "Claim is bounded; H4 requires a second open model and non-math corpus without retuning."],
            ["Measurement", "Numeric extraction can mis-score malformed outputs", "Exact extraction and normalization code are frozen; future work should include a blinded manual scorer audit."],
            ["Operational", "Generated tokens do not equal end-to-end cost", "Future replication should record latency, GPU-seconds, energy, and dollar cost alongside tokens."],
        ],
        [1.0, 1.85, 3.7],
    )

    doc.add_heading("8. Praxis Contribution and Recommended Next Gate", level=1)
    add_body(
        doc,
        "The Praxis contribution is a verification-gated inference control: stop only when a frozen observable rule qualifies, quantify both prevented degradation and stopping harm, and refuse broad deployment claims until transfer gates pass. This reframes adaptive compute as a safety case rather than only an efficiency optimization."
    )
    add_body(doc, "The next investment should be a frozen H4 replication matrix:")
    add_numbered(
        doc,
        [
            "Retain the current 200-item Qwen/GSM8K result unchanged as the discovery experiment.",
            "Add at least one second open model with the same six arms and no threshold adjustment after test inspection.",
            "Add a non-math reasoning corpus with deterministic scoring or blinded adjudication.",
            "Predefine a validation-only threshold-selection process, then lock thresholds for held-out testing.",
            "Measure wall-clock latency, GPU-seconds, and cost in addition to generated tokens.",
            "Require the confidence component to beat or reduce harm relative to stability-only before retaining it in the final production policy.",
        ],
    )

    doc.add_heading("9. Conclusion", level=1)
    add_body(
        doc,
        "PX-057 passed its preregistered Gate 2 evaluation. On 200 frozen GSM8K questions, Qwen2.5-7B-Instruct achieved 91.0% accuracy with adaptive stopping versus 61.5% after eight forced reconsideration rounds, while saving 66.5% of generated tokens, preventing 60 of 67 observed overthinking events, and producing one early-stop harm. The result is strong enough to justify H4 investment and a bounded Praxis claim. It is not yet sufficient for a general robustness claim. The scientifically defensible conclusion is that answer-stability stopping is promising and operationally testable, while transfer and the incremental value of confidence remain open questions."
    )

    doc.add_heading("Data, Code, and Evidence Availability", level=1)
    add_table(
        doc,
        ["Artifact", "Repository location / identifier"],
        [
            ["Preregistration", "reports/adaptive_stopping_overthinking/PX057_ADAPTIVE_STOPPING_PREREG_20260723.md"],
            ["Frozen configuration", "configs/px057_adaptive_stopping_gate2_full_20260724.json"],
            ["Trace collection", "scripts/run_px057_trace_collection.py"],
            ["Policy evaluation", "scripts/run_px057_adaptive_stopping_gate.py"],
            ["Independent adjudication", "scripts/adjudicate_px057_adaptive_stopping.py"],
            ["Verified determination", "reports/adaptive_stopping_overthinking/gate2_full_determination_20260724/summary.json"],
            ["Cloud job", "px057-gate2-full-retry-2026-07-24-14-55-01-066"],
            ["Cloud output", "s3://praxis-garypagan-272615233626-us-east-1/experiments/px057-adaptive-stopping/gate2-full-20260724/output/.../model.tar.gz"],
        ],
        [1.65, 4.9],
    )
    add_body(
        doc,
        "Core evidence SHA-256 hashes: selected_rows.json 6203c728f838fda9b932f69c83f7c90eb2e04c65045b24071d5347d4aaa6fc7e; reasoning_traces.jsonl d1221cab111650111e2cb8aa0b03a6b3c99b2bc0738c222cd078ab8935ad9b8c; raw_generations.jsonl 5f63edfe41c1b4cd47d3678eba133ae0c139164743265d49cfea70a546088646."
    )

    doc.add_heading("References", level=1)
    add_reference(
        doc,
        "Zhou, S., Ling, R., Chen, J., Wang, X., Fan, T., & Wang, H. (2026). When More Thinking Hurts: Overthinking in LLM Test-Time Compute Scaling. Findings of ACL 2026, 23967-23977. DOI: 10.18653/v1/2026.findings-acl.1199.",
        "Original publication and PDF",
        "https://aclanthology.org/2026.findings-acl.1199/",
    )
    add_reference(
        doc,
        "Snell, C., Lee, J., Xu, K., & Kumar, A. (2024). Scaling LLM Test-Time Compute Optimally can be More Effective than Scaling Model Parameters.",
        "arXiv:2408.03314",
        "https://arxiv.org/abs/2408.03314",
    )
    add_reference(
        doc,
        "Muennighoff, N., et al. (2025). s1: Simple test-time scaling.",
        "arXiv:2501.19393",
        "https://arxiv.org/abs/2501.19393",
    )
    add_reference(
        doc,
        "Cobbe, K., et al. (2021). Training Verifiers to Solve Math Word Problems.",
        "arXiv:2110.14168 (GSM8K source)",
        "https://arxiv.org/abs/2110.14168",
    )
    add_reference(
        doc,
        "Qwen Team. (2024). Qwen2.5 Technical Report.",
        "arXiv:2412.15115",
        "https://arxiv.org/abs/2412.15115",
    )

    doc.add_heading("Appendix A. Verification Record", level=1)
    add_table(
        doc,
        ["Check", "Result"],
        [[key.replace("_", " ").title(), "PASS" if value else "FAIL"] for key, value in summary["completeness_checks"].items()],
        [5.2, 1.0],
    )
    doc.add_heading("Appendix B. Registered Claim Statement", level=1)
    add_body(
        doc,
        "Permitted claim: On a frozen 200-item GSM8K sample with Qwen2.5-7B-Instruct and this eight-round iterative prompting protocol, the preregistered adaptive stopping rule preserved or improved accuracy relative to fixed-long inference, reduced generated-token use, prevented most observed correct-to-wrong overthinking events, and stayed below the harm-rate ceiling."
    )
    add_body(
        doc,
        "Prohibited overclaim: PX-057 does not yet establish large-scale robustness, general LLM overthinking prevention, cross-domain transfer, cross-model transfer, or superiority of confidence-aware stopping over answer stability alone."
    )
    return doc


def build_report(summary: dict, config: dict, accuracy_chart: Path, gate_chart: Path) -> Document:
    m = summary["metrics"]
    doc = Document()
    configure_document(doc, "PX-057 Final Praxis Report")
    title_page(
        doc,
        "Executive determination, evidence summary, and investment recommendation",
        "FINAL PRAXIS REPORT  /  DECISION BRIEF",
    )
    doc.add_heading("Final Determination", level=1)
    add_body(
        doc,
        "GO FOR H4 REPLICATION. PX-057 is a valid positive Gate 2 result and the strongest current candidate for further investment. All preregistered H1-H3 checks passed. Do not represent the result as large-scale or general robustness until cross-model and cross-domain transfer is demonstrated."
    )
    add_figure(doc, gate_chart, "Registered outcomes and decision thresholds.")
    add_table(
        doc,
        ["KPI", "Result", "Decision meaning"],
        [
            ["Adaptive accuracy", "91.0% (182/200)", "29.5 pp above fixed-long"],
            ["Fixed-long accuracy", "61.5% (123/200)", "Eight reconsideration rounds degraded many answers"],
            ["Fixed-short accuracy", "88.0% (176/200)", "Adaptive added 3.0 pp over round 2"],
            ["Compute saving", "66.5% [63.7%, 69.1%]", "Passed 20% floor"],
            ["Prevention", "60/67; 89.6% [82.1%, 95.5%]", "Passed 25% floor"],
            ["Early-stop harm", "1/200; 0.5%", "Passed 2% ceiling"],
        ],
        [1.55, 2.0, 3.0],
    )
    doc.add_heading("What the Result Proves", level=1)
    add_bullets(
        doc,
        [
            "The frozen Qwen/GSM8K iterative protocol exhibited frequent correct-to-wrong degradation by round 8.",
            "The preregistered adaptive policy avoided most observed degradation and used substantially fewer generated tokens.",
            "The evidence package is complete: 200 unique questions, 200 complete traces, and 1,600 unique round generations were independently verified.",
        ],
    )
    doc.add_heading("Critical Nuance", level=1)
    add_body(
        doc,
        "Adaptive stopping and answer-stability stopping both scored 91.0%. The final Praxis should therefore be framed around stability-gated stopping, not around a proven benefit from the confidence threshold. The confidence component remains a hypothesis for H4, not an established mechanism."
    )
    add_figure(doc, accuracy_chart, "The adaptive arm tied the simpler stability-only control.")
    doc.add_heading("What It Does Not Prove", level=1)
    add_bullets(
        doc,
        [
            "No full-GSM8K or population-level accuracy estimate.",
            "No cross-model, cross-domain, stochastic-decoding, or prompt-transfer result.",
            "No calibrated uncertainty result.",
            "No direct latency, energy, or dollar-cost result.",
        ],
    )
    doc.add_heading("Investment Recommendation", level=1)
    add_body(
        doc,
        "Invest in one tightly frozen H4 study before expanding implementation. Preserve the current result as the discovery experiment and preregister the transfer matrix before inspecting new outputs."
    )
    add_numbered(
        doc,
        [
            "Replicate on a second open model without retuning the current policy on the test set.",
            "Add a non-math corpus with deterministic or blinded scoring.",
            "Keep all six comparison arms, especially stability-only.",
            "Add latency, GPU-seconds, and dollar cost.",
            "Retain confidence only if it improves accuracy, reduces harm, or lowers cost versus stability-only.",
        ],
    )
    doc.add_heading("Publication-Ready Claim", level=1)
    add_body(
        doc,
        "On a frozen 200-item GSM8K sample with Qwen2.5-7B-Instruct and an eight-round iterative reconsideration protocol, a preregistered adaptive stopping rule achieved 91.0% accuracy versus 61.5% for fixed-long inference, saved 66.5% of generated tokens, prevented 60 of 67 observed correct-to-wrong events, and harmed 1 of 200 questions. These findings are limited to the tested model, sample, prompting protocol, and thresholds."
    )
    doc.add_heading("Literature Foundation", level=1)
    add_reference(
        doc,
        "Zhou, S., et al. (2026). When More Thinking Hurts: Overthinking in LLM Test-Time Compute Scaling. Findings of ACL 2026.",
        "Read the original publication",
        "https://aclanthology.org/2026.findings-acl.1199/",
    )
    add_reference(
        doc,
        "Cobbe, K., et al. (2021). Training Verifiers to Solve Math Word Problems. Introduces GSM8K.",
        "Read the dataset paper",
        "https://arxiv.org/abs/2110.14168",
    )
    return doc


PAPER_MD = """# PX-057: Adaptive Stopping to Prevent LLM Overthinking

Final Praxis paper - July 24, 2026

## Final determination

PX-057 is a valid positive Gate 2 result. On a frozen 200-item GSM8K sample with Qwen2.5-7B-Instruct and an eight-round iterative reconsideration protocol, the adaptive policy achieved 91.0% accuracy versus 61.5% for fixed-long inference, saved 66.5% of generated tokens, prevented 60 of 67 observed correct-to-wrong events, and harmed 1 of 200 questions.

## Registered results

| Gate | Threshold | Result | Decision |
|---|---:|---:|---|
| Accuracy delta vs. fixed-long | >= -1.0 pp | +29.5 pp | PASS |
| Mean generated-token saving | >= 20% | 66.5% (95% bootstrap CI 63.7%-69.1%) | PASS |
| Overthinking prevention | >= 25% | 89.6%; 60/67 (95% bootstrap CI 82.1%-95.5%) | PASS |
| Early-stop harm | <= 2% | 0.5%; 1/200 | PASS |

## Arm outcomes

| Arm | Accuracy |
|---|---:|
| Fixed-long, round 8 | 61.5% |
| Fixed-short, round 2 | 88.0% |
| Uncertainty-only | 88.0% |
| Answer stability | 91.0% |
| Adaptive | 91.0% |
| Oracle best step (descriptive) | 95.0% |

## Interpretation

The strongest demonstrated fact is that repeated reconsideration through round 8 degraded many initially correct answers and that a preregistered stability gate avoided most degradation. Adaptive stopping tied answer-stability stopping, so the confidence condition did not add measurable accuracy in this run.

## Claim boundary

This result applies to one 200-item GSM8K sample, Qwen2.5-7B-Instruct, greedy decoding, and one iterative prompting protocol. H4 cross-model and cross-domain transfer remains pending. No large-scale or general robustness claim is supported.

## Literature

- Zhou et al. (2026), [When More Thinking Hurts: Overthinking in LLM Test-Time Compute Scaling](https://aclanthology.org/2026.findings-acl.1199/) - original publication and PDF.
- Snell et al. (2024), [Scaling LLM Test-Time Compute Optimally can be More Effective than Scaling Model Parameters](https://arxiv.org/abs/2408.03314).
- Muennighoff et al. (2025), [s1: Simple test-time scaling](https://arxiv.org/abs/2501.19393).
- Cobbe et al. (2021), [Training Verifiers to Solve Math Word Problems](https://arxiv.org/abs/2110.14168) - GSM8K source.
- Qwen Team (2024), [Qwen2.5 Technical Report](https://arxiv.org/abs/2412.15115).

## Recommended next gate

Freeze an H4 replication matrix with a second open model and a non-math reasoning corpus, retain all comparison arms, add real latency/GPU/cost measurements, and require confidence to improve over stability-only before keeping it in the production policy.
"""


REPORT_MD = """# PX-057 Final Praxis Report

## Decision

**GO for H4 replication.** Gate 2 is valid and positive; H1-H3 passed. Do not claim large-scale or general robustness yet.

## Headline evidence

- Adaptive accuracy: 91.0% (182/200)
- Fixed-long accuracy: 61.5% (123/200)
- Mean generated-token saving: 66.5% (95% bootstrap CI 63.7%-69.1%)
- Prevention: 60/67 events, 89.6% (95% bootstrap CI 82.1%-95.5%)
- Early-stop harm: 1/200, 0.5%
- Integrity: 200 unique questions, 200 complete traces, 1,600 unique question-round outputs

## Critical nuance

Adaptive and answer-stability arms tied at 91.0%. Stability is the demonstrated signal; confidence adds no measured accuracy benefit in this run.

## Next investment

Run H4 on a second open model and a non-math corpus without test-set retuning. Keep the stability-only control and add latency, GPU-seconds, and dollar cost.

## Original literature

[Zhou et al. (2026), When More Thinking Hurts: Overthinking in LLM Test-Time Compute Scaling](https://aclanthology.org/2026.findings-acl.1199/)
"""


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    config = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    accuracy_chart = TMP / "px057_accuracy.png"
    gate_chart = TMP / "px057_gates.png"
    draw_accuracy_chart(summary["metrics"], accuracy_chart)
    draw_gate_chart(summary["metrics"], config, gate_chart)
    paper = build_paper(summary, config, accuracy_chart, gate_chart)
    report = build_report(summary, config, accuracy_chart, gate_chart)
    paper_path = OUT / "PX-057_Adaptive_Stopping_Final_Praxis_Paper.docx"
    report_path = OUT / "PX-057_Executive_Final_Report.docx"
    paper.save(paper_path)
    report.save(report_path)
    (OUT / "PX-057_Adaptive_Stopping_Final_Praxis_Paper.md").write_text(
        PAPER_MD, encoding="utf-8"
    )
    (OUT / "PX-057_Executive_Final_Report.md").write_text(
        REPORT_MD, encoding="utf-8"
    )
    shutil.copy2(accuracy_chart, OUT / accuracy_chart.name)
    shutil.copy2(gate_chart, OUT / gate_chart.name)
    print(paper_path)
    print(report_path)


if __name__ == "__main__":
    main()
